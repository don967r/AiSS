import streamlit as st
import pandas as pd
import geopandas as gpd
import folium
from folium.plugins import HeatMap
from streamlit_folium import st_folium
from datetime import timedelta, datetime
import io
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt # <-- Новый импорт

# --- 1. Конфигурация страницы и Заголовок ---
st.set_page_config(layout="wide", page_title="Анализ 'Судно-Пятно'")

# --- CSS для точной настройки вертикальных отступов ---
st.markdown("""
<style>
/* Уменьшаем общие вертикальные отступы для всей страницы */
div.block-container {
    padding-top: 2rem;
    padding-bottom: 1rem;
}
/* Уменьшаем стандартный промежуток между элементами */
div[data-testid="stVerticalBlock"] {
    gap: 0.8rem;
}
/* Убираем лишний нижний отступ у контейнера с картой */
div[data-testid="stFolium"] {
    margin-bottom: 0 !important;
}
/* Настраиваем отступы у заголовков для лучшего вида */
h2 {
    margin-top: 1.5rem !important;
    margin-bottom: 0.5rem !important;
}
/* Заставляет панель вкладки сжиматься до высоты своего контента */
div[data-testid="stTabsPanel"] {
    padding-top: 1rem;
    padding-bottom: 0 !important;
    min-height: 0;
}
</style>
""", unsafe_allow_html=True)

# --- Задаем пути к файлам в репозитории ---
SPILLS_FILE_PATH = 'fields2.geojson'
AIS_FILE_PATH = 'generated_ais_data.csv'
ROUTES_FILE_PATH = 'routs.geojson'

# --- 2. Боковая панель с параметрами ---
st.sidebar.header("Параметры анализа")

time_window_hours = st.sidebar.slider(
    "Временное окно поиска (часы до обнаружения):",
    min_value=1, max_value=168, value=24, step=1,
    help="Искать суда, которые были в зоне разлива за указанное количество часов ДО его обнаружения."
)
date_range = st.sidebar.date_input(
    "Диапазон дат для анализа:",
    value=(datetime(2022, 1, 1), datetime(2025, 7, 15)),
    min_value=datetime(2000, 1, 1),
    max_value=datetime(2030, 12, 31),
    help="Выберите диапазон дат для фильтрации разливов и AIS-данных."
)
st.sidebar.header("Управление слоями")
show_spills = st.sidebar.checkbox("Пятна разливов", value=True)
show_ships = st.sidebar.checkbox("Суда-кандидаты", value=True)
show_routes = st.sidebar.checkbox("Судовые трассы", value=True)

# --- 3. Функции для обработки и анализа данных ---
@st.cache_data
def load_spills_data(file_path):
    try:
        gdf = gpd.read_file(file_path)
    except Exception as e:
        st.error(f"Не удалось прочитать GeoJSON файл '{file_path}'. Ошибка: {e}")
        return gpd.GeoDataFrame()
    required_cols = ['slick_name', 'area_sys']
    if not all(col in gdf.columns for col in required_cols):
        missing = [col for col in required_cols if col not in gdf.columns]
        st.error(f"В GeoJSON отсутствуют обязательные поля: {', '.join(missing)}")
        return gpd.GeoDataFrame()
    gdf.rename(columns={'slick_name': 'spill_id', 'area_sys': 'area_sq_km'}, inplace=True)
    if 'date' in gdf.columns and 'time' in gdf.columns:
        gdf['detection_date'] = pd.to_datetime(gdf['date'] + ' ' + gdf['time'], errors='coerce')
    else:
        gdf['detection_date'] = pd.to_datetime(gdf['spill_id'], format='%Y-%m-%d_%H:%M:%S', errors='coerce')
    gdf.dropna(subset=['detection_date'], inplace=True)
    if gdf.crs is None:
        gdf.set_crs("EPSG:4326", inplace=True)
    else:
        gdf = gdf.to_crs("EPSG:4326")
    return gdf

@st.cache_data
def load_ais_data(file_path):
    try:
        df = pd.read_csv(file_path, low_memory=False)
    except Exception as e:
        st.error(f"Не удалось прочитать CSV файл '{file_path}'. Ошибка: {e}")
        return gpd.GeoDataFrame()
    required_cols = ['mmsi', 'latitude', 'longitude', 'BaseDateTime']
    if not all(col in df.columns for col in required_cols):
        missing = [col for col in required_cols if col not in df.columns]
        st.error(f"В CSV файле отсутствуют обязательные колонки: {', '.join(missing)}")
        return gpd.GeoDataFrame()
    df['timestamp'] = pd.to_datetime(df['BaseDateTime'], errors='coerce')
    df.dropna(subset=['timestamp', 'latitude', 'longitude'], inplace=True)
    gdf = gpd.GeoDataFrame(
        df,
        geometry=gpd.points_from_xy(df.longitude, df.latitude),
        crs="EPSG:4326"
    )
    return gdf

@st.cache_data
def load_routes_data(file_path):
    try:
        gdf = gpd.read_file(file_path)
        if gdf.crs is None:
            gdf.set_crs("EPSG:4326", inplace=True)
        else:
            gdf = gdf.to_crs("EPSG:4326")
        return gdf
    except Exception as e:
        st.warning(f"Не удалось прочитать файл трасс '{file_path}'. Слой не будет отображен. Ошибка: {e}")
        return gpd.GeoDataFrame()

def find_candidates(spills_gdf, vessels_gdf, time_window_hours):
    if spills_gdf.empty or vessels_gdf.empty:
        return gpd.GeoDataFrame()
    candidates = gpd.sjoin(vessels_gdf, spills_gdf, predicate='within')
    if candidates.empty:
        return gpd.GeoDataFrame()
    time_delta = timedelta(hours=time_window_hours)
    candidates = candidates[
        (candidates['timestamp'] <= candidates['detection_date']) &
        (candidates['timestamp'] >= candidates['detection_date'] - time_delta)
    ]
    return candidates

# --- ФУНКЦИИ ДЛЯ ГЕНЕРАЦИИ ОТЧЕТА ---

# >>> НОВАЯ ФУНКЦИЯ ДЛЯ ГЕНЕРАЦИИ ГРАФИКА
def create_incident_plot(spill_data, candidates_data):
    """
    Создает график инцидента (пятно + суда) с помощью Matplotlib.
    Возвращает изображение в виде байтового потока.
    """
    fig, ax = plt.subplots(figsize=(10, 8))

    # Отрисовка контура пятна
    gpd.GeoSeries([spill_data['geometry']]).plot(
        ax=ax,
        edgecolor='red',
        facecolor='red',
        alpha=0.3,
        linewidth=1.5,
        label=f"Пятно {spill_data['spill_id']}"
    )

    # Отрисовка судов-кандидатов
    if not candidates_data.empty:
        candidates_data.plot(
            ax=ax,
            marker='o',
            color='blue',
            markersize=50,
            label='Суда-кандидаты'
        )
        # Добавляем аннотации MMSI к точкам судов
        for idx, row in candidates_data.iterrows():
            ax.text(row.geometry.x, row.geometry.y, f" {row['mmsi']}", fontsize=9, ha='left', color='navy')


    # Настройка графика
    ax.set_title(f"Карта-схема инцидента: {spill_data['spill_id']}", fontsize=14)
    ax.set_xlabel("Долгота")
    ax.set_ylabel("Широта")
    ax.grid(True, linestyle='--', alpha=0.6)
    ax.legend()
    # Устанавливаем равный масштаб для осей, чтобы избежать искажений
    ax.set_aspect('equal', adjustable='box')
    plt.tight_layout()

    # Сохраняем график в байтовый буфер в памяти
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150)
    buf.seek(0)
    plt.close(fig) # Закрываем фигуру, чтобы освободить память

    return buf.getvalue()

def strfdelta(tdelta, fmt):
    """Функция для красивого форматирования timedelta."""
    d = {"days": tdelta.days}
    d["hours"], rem = divmod(tdelta.seconds, 3600)
    d["minutes"], d["seconds"] = divmod(rem, 60)
    return fmt.format(**d)

def generate_docx_report(spill_data, candidates_data, prime_suspect_data, historical_data, plot_bytes=None):
    """Генерирует отчет в формате DOCX, включая опциональный график."""
    doc = Document()
    doc.add_heading('ОТЧЕТ ПО АНАЛИЗУ СВЯЗИ "СУДНО-ПЯТНО"', level=1)
    
    doc.add_paragraph(f"Дата составления отчета: {datetime.now().strftime('%d.%m.%Y')}")
    doc.add_paragraph(f"ID отчета: SP-{spill_data['spill_id'].replace(':', '-')}")
    doc.add_paragraph()

    doc.add_heading('1. ИНФОРМАЦИЯ ОБ ИНЦИДЕНТЕ (ПЯТНЕ)', level=2)
    doc.add_paragraph(f"ID пятна (spill_id): {spill_data['spill_id']}")
    doc.add_paragraph(f"Дата и время обнаружения: {spill_data['detection_date'].strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Площадь пятна (км²): {spill_data.get('area_sq_km', 0):.2f}")
    
    centroid = spill_data['geometry'].centroid
    doc.add_paragraph(f"Географические координаты (центроид): Lat {centroid.y:.4f}, Lon {centroid.x:.4f}")

    if plot_bytes:
        try:
            image_stream = io.BytesIO(plot_bytes)
            doc.add_picture(image_stream, width=Inches(6.0))
            p = doc.add_paragraph(f"Рис. 1. Карта-схема расположения пятна {spill_data['spill_id']} и судов-кандидатов.")
            p.style = 'Caption'
        except Exception as e:
            doc.add_paragraph(f"[Не удалось вставить изображение: {e}]")
    
    doc.add_heading('2. ПАРАМЕТРЫ АНАЛИЗА', level=2)
    doc.add_paragraph(f"Временное окно поиска: {time_window_hours} часов до момента обнаружения пятна.")
    doc.add_paragraph("Критерий связи: Судно считается кандидатом, если его AIS-позиция зафиксирована внутри географических границ пятна в указанном временном окне.")

    doc.add_heading('3. РЕЗУЛЬТАТЫ АНАЛИЗА: СУДА-КАНДИДАТЫ', level=2)
    if not candidates_data.empty:
        doc.add_paragraph("Сводная таблица судов-кандидатов:")
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'MMSI судна'
        hdr_cells[1].text = 'Название судна'
        hdr_cells[2].text = 'Время прохода'
        hdr_cells[3].text = 'Время до обнаружения'

        for _, row in candidates_data.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row.get('mmsi', 'N/A'))
            row_cells[1].text = str(row.get('vessel_name', 'Имя не указано'))
            row_cells[2].text = row['timestamp'].strftime('%Y-%m-%d %H:%M')
            time_to_detection = spill_data['detection_date'] - row['timestamp']
            row_cells[3].text = strfdelta(time_to_detection, "{hours} ч {minutes} мин")

        doc.add_paragraph()
        doc.add_heading('Детальный анализ основного кандидата', level=3)
        if prime_suspect_data is not None and not prime_suspect_data.empty:
            suspect = prime_suspect_data.iloc[0]
            doc.add_paragraph(f"Судно: {suspect.get('vessel_name', 'Имя не указано')} (MMSI: {suspect['mmsi']})")
            doc.add_paragraph(f"Время до обнаружения: {strfdelta(suspect['time_to_detection'], '{hours} ч {minutes} мин')}")
            if historical_data:
                 doc.add_paragraph(f"Исторический контекст: {historical_data.get('incident_count', 0)} связанных инцидентов, общая площадь {historical_data.get('total_area_sq_km', 0):.2f} км².")
        else:
            doc.add_paragraph("Основной кандидат (с минимальным временем до обнаружения) не определен.")
    else:
        doc.add_paragraph("Суда-кандидаты в заданном временном окне не найдены.")

    doc.add_heading('4. ЗАКЛЮЧЕНИЕ', level=2)
    conclusion_text = (
        f"1. Факт: {spill_data['detection_date'].strftime('%d.%m.%Y')} было зафиксировано загрязнение (ID: {spill_data['spill_id']}) "
        f"площадью {spill_data.get('area_sq_km', 0):.2f} км².\n"
        f"2. Анализ: В результате анализа данных AIS за {time_window_hours} часов до обнаружения было выявлено {len(candidates_data)} судов-кандидатов.\n"
    )
    if prime_suspect_data is not None and not prime_suspect_data.empty:
        suspect = prime_suspect_data.iloc[0]
        conclusion_text += (
            f"3. Основной вывод: Наиболее вероятным источником загрязнения является судно "
            f"{suspect.get('vessel_name', 'Имя не указано')} (MMSI: {suspect['mmsi']}), так как оно прошло через данную точку за "
            f"{strfdelta(suspect['time_to_detection'], '{hours} ч {minutes} мин')} до фиксации пятна."
        )
    else:
        conclusion_text += "3. Основной вывод: Определить наиболее вероятный источник не удалось из-за отсутствия судов-кандидатов."
    doc.add_paragraph(conclusion_text)

    doc.add_heading('5. РЕКОМЕНДАЦИИ', level=2)
    if prime_suspect_data is not None and not prime_suspect_data.empty:
        suspect = prime_suspect_data.iloc[0]
        reco_text = (
            f"- Провести углубленную проверку в отношении судна {suspect.get('vessel_name', 'Имя не указано')} (MMSI: {suspect['mmsi']}).\n"
            "- Запросить судовые документы и записи бортовых журналов за период, предшествующий инциденту."
        )
        doc.add_paragraph(reco_text)
    else:
        doc.add_paragraph("Рекомендации не могут быть сформированы из-за отсутствия основного кандидата.")
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.getvalue()

# --- 4. Основная логика приложения ---
spills_gdf = load_spills_data(SPILLS_FILE_PATH)
vessels_gdf = load_ais_data(AIS_FILE_PATH)
routes_gdf = load_routes_data(ROUTES_FILE_PATH)

if spills_gdf.empty or vessels_gdf.empty:
    st.error("Не удалось загрузить или обработать основные файлы данных (разливы, AIS). Анализ остановлен.")
    st.stop()

if len(date_range) == 2:
    start_date, end_date = pd.to_datetime(date_range[0]), pd.to_datetime(date_range[1]) + pd.Timedelta(days=1)
    spills_gdf = spills_gdf[(spills_gdf['detection_date'] >= start_date) & (spills_gdf['detection_date'] <= end_date)]
    vessels_gdf = vessels_gdf[(vessels_gdf['timestamp'] >= start_date) & (vessels_gdf['timestamp'] <= end_date)]

vessel_options = vessels_gdf[['mmsi']].drop_duplicates()
if 'vessel_name' in vessels_gdf.columns:
    vessel_options = vessels_gdf[['mmsi', 'vessel_name']].drop_duplicates()
    vessel_options['display'] = vessel_options.apply(
        lambda x: f"{x['vessel_name']} (MMSI: {x['mmsi']})" if pd.notnull(x['vessel_name']) else f"MMSI: {x['mmsi']}", axis=1
    )
else:
    vessel_options['display'] = vessel_options['mmsi'].apply(lambda x: f"MMSI: {x}")

selected_vessels_display = st.sidebar.multiselect(
    "Выберите суда для фильтрации:",
    options=vessel_options['display'].tolist(),
    help="Фильтрует данные на карте и в таблицах по выбранным судам."
)

filtered_routes_gdf = routes_gdf.copy()
if selected_vessels_display:
    selected_mmsi = vessel_options[vessel_options['display'].isin(selected_vessels_display)]['mmsi'].tolist()
    vessels_gdf = vessels_gdf[vessels_gdf['mmsi'].isin(selected_mmsi)]
    if not filtered_routes_gdf.empty and 'mmsi' in filtered_routes_gdf.columns:
        filtered_routes_gdf = filtered_routes_gdf[filtered_routes_gdf['mmsi'].isin(selected_mmsi)]

# --- 5. Отображение карты и таблицы ---
with st.container(border=False):
    st.header("Карта разливов и судов-кандидатов")
    
    map_center = [67.63778, 53.00667] 
    map_tiles = "CartoDB dark_matter"
    m = folium.Map(location=map_center, zoom_start=3, tiles=map_tiles, attributionControl=False)
    
    candidates_df = gpd.GeoDataFrame()
    if spills_gdf.empty:
        st.warning("Нет данных о разливах в выбранном диапазоне дат.")
    else:
        candidates_df = find_candidates(spills_gdf, vessels_gdf, time_window_hours)
        spills_fg = folium.FeatureGroup(name="Пятна разливов", show=show_spills)
        for _, row in spills_gdf.iterrows():
            folium.GeoJson(
                row['geometry'],
                style_function=lambda x: {'fillColor': '#B22222', 'color': 'black', 'weight': 1.5, 'fillOpacity': 0.6},
                tooltip=f"<b>Пятно:</b> {row.get('spill_id', 'N/A')}<br>"
                        f"<b>Время:</b> {row['detection_date'].strftime('%Y-%m-%d %H:%M')}<br>"
                        f"<b>Площадь:</b> {row.get('area_sq_km', 0):.2f} км²"
            ).add_to(spills_fg)
        spills_fg.add_to(m)

        candidate_vessels_fg = folium.FeatureGroup(name="Суда-кандидаты", show=show_ships)
        if not candidates_df.empty:
            for _, row in candidates_df.iterrows():
                vessel_name = row.get('vessel_name', 'Имя не указано')
                folium.Marker(
                    location=[row.geometry.y, row.geometry.x],
                    tooltip=f"<b>Судно:</b> {vessel_name} (MMSI: {row['mmsi']})<br>"
                            f"<b>Время прохода:</b> {row['timestamp'].strftime('%Y-%m-%d %H:%M')}<br>"
                            f"<b>Внутри пятна:</b> {row['spill_id']}",
                    icon=folium.Icon(color='blue', icon='ship', prefix='fa')
                ).add_to(candidate_vessels_fg)
        candidate_vessels_fg.add_to(m)
        
        routes_fg = folium.FeatureGroup(name="Судовые трассы", show=show_routes)
        if not filtered_routes_gdf.empty:
            for _, row in filtered_routes_gdf.iterrows():
                tooltip_text = f"<b>Трек судна (MMSI: {row.get('mmsi', 'N/A')})</b>"
                folium.GeoJson(
                    row['geometry'],
                    style_function=lambda x: {'color': '#007FFF', 'weight': 2.5, 'opacity': 0.7},
                    tooltip=tooltip_text
                ).add_to(routes_fg)
        routes_fg.add_to(m)

    folium.LayerControl().add_to(m) 
    st_folium(m, width=1200, height=400, returned_objects=[])

    st.header(f"Таблица судов-кандидатов (в пределах {time_window_hours} часов)")
    if candidates_df.empty:
        st.info("В заданном временном окне и с учетом фильтров суда-кандидаты не найдены.")
    else:
        report_df = candidates_df.drop_duplicates(subset=['spill_id', 'mmsi'])
        desired_cols = ['spill_id', 'mmsi', 'vessel_name', 'timestamp', 'detection_date', 'area_sq_km']
        existing_cols = [col for col in desired_cols if col in report_df.columns]
        display_df = report_df[existing_cols].copy()
        rename_dict = {
            'spill_id': 'ID Пятна', 'mmsi': 'MMSI Судна', 'vessel_name': 'Название судна',
            'timestamp': 'Время прохода', 'detection_date': 'Время обнаружения', 'area_sq_km': 'Площадь, км²'
        }
        display_df.rename(columns=rename_dict, inplace=True)
        st.dataframe(display_df.sort_values(by='Время обнаружения', ascending=False).reset_index(drop=True), use_container_width=True)

# --- 6. Блок с расширенной аналитикой ---
with st.container(border=False):
    st.header("Дополнительная аналитика")
    tab1, tab2, tab3 = st.tabs(["📊 Аналитика по судам", "📍 Горячие точки (Hotspots)", "🔍 Аналитика по инцидентам"])

    candidates_df_for_analytics = find_candidates(spills_gdf, vessels_gdf, time_window_hours) if not spills_gdf.empty else gpd.GeoDataFrame()
    prime_suspects_df = pd.DataFrame()

    with tab1:
        if not candidates_df_for_analytics.empty:
            unique_incidents = candidates_df_for_analytics.drop_duplicates(subset=['mmsi', 'spill_id'])
            ship_names = unique_incidents[['mmsi', 'vessel_name']].drop_duplicates('mmsi')
            st.subheader("Антирейтинг по количеству связанных пятен")
            ship_incident_counts = unique_incidents.groupby('mmsi').size().reset_index(name='incident_count').sort_values('incident_count', ascending=False)
            st.dataframe(pd.merge(ship_incident_counts, ship_names, on='mmsi', how='left'), use_container_width=True)
            st.subheader("Антирейтинг по суммарной площади связанных пятен (км²)")
            ship_area_sum = unique_incidents.groupby('mmsi')['area_sq_km'].sum().reset_index(name='total_area_sq_km').sort_values('total_area_sq_km', ascending=False)
            st.dataframe(pd.merge(ship_area_sum, ship_names, on='mmsi', how='left'), use_container_width=True)
        else:
            st.info("Нет данных для аналитики по судам.")

    with tab2:
        st.subheader("Карта 'горячих точек' разливов")
        if spills_gdf.empty:
            st.warning("Нет данных для отображения карты горячих точек.")
        else:
            map_center = [67.638, 53.005]
            m_heatmap = folium.Map(location=map_center, zoom_start=3, tiles=map_tiles, attributionControl=False)
            heat_data = [[point.xy[1][0], point.xy[0][0], row['area_sq_km']] for _, row in spills_gdf.iterrows() for point in [row['geometry'].centroid]]
            HeatMap(heat_data, radius=15, blur=20).add_to(m_heatmap)
            st_folium(m_heatmap, width=1200, height=350, returned_objects=[])

    with tab3:
        if not candidates_df_for_analytics.empty:
            unique_incidents = candidates_df_for_analytics.drop_duplicates(subset=['mmsi', 'spill_id'])
            st.subheader("Пятна с наибольшим количеством судов-кандидатов")
            spill_candidate_counts = candidates_df_for_analytics.groupby('spill_id')['mmsi'].nunique().reset_index(name='candidate_count').sort_values('candidate_count', ascending=False)
            st.dataframe(spill_candidate_counts, use_container_width=True)
            
            st.subheader("Главные подозреваемые (минимальное время до обнаружения)")
            candidates_df_for_analytics['time_to_detection'] = candidates_df_for_analytics['detection_date'] - candidates_df_for_analytics['timestamp']
            prime_suspects_idx = candidates_df_for_analytics.groupby('spill_id')['time_to_detection'].idxmin()
            prime_suspects_df = candidates_df_for_analytics.loc[prime_suspects_idx]
            display_cols = ['spill_id', 'mmsi', 'vessel_name', 'time_to_detection', 'area_sq_km']
            st.dataframe(prime_suspects_df[[col for col in display_cols if col in prime_suspects_df]].sort_values('area_sq_km', ascending=False), use_container_width=True)
        else:
            st.info("Нет данных для аналитики по инцидентам.")

# --- 7. УЛУЧШЕННЫЙ БЛОК ФОРМИРОВАНИЯ ОТЧЕТА С АВТОМАТИЧЕСКИМ ГРАФИКОМ ---
st.header("Формирование отчета по инциденту")
with st.expander("🖨️ Нажмите, чтобы выбрать инцидент и создать отчет", expanded=False):

    if not candidates_df_for_analytics.empty:
        # --- Часть 1: Выбор инцидента для отчета ---
        spills_with_candidates = candidates_df_for_analytics[['spill_id', 'detection_date', 'area_sq_km']].drop_duplicates(subset=['spill_id'])
        candidate_counts = candidates_df_for_analytics.groupby('spill_id')['mmsi'].nunique().reset_index(name='candidate_count')
        
        reportable_incidents_df = pd.merge(spills_with_candidates, candidate_counts, on='spill_id').sort_values(by='detection_date', ascending=False)
        reportable_incidents_df.rename(columns={
            'spill_id': 'ID Пятна', 'detection_date': 'Дата обнаружения',
            'area_sq_km': 'Площадь, км²', 'candidate_count': 'Кол-во канд.'
        }, inplace=True)

        st.info("Ниже представлены инциденты, для которых найдены суда-кандидаты. Выберите один для создания отчета.")
        st.dataframe(reportable_incidents_df.reset_index(drop=True), use_container_width=True)

        reportable_incidents_df['display_option'] = reportable_incidents_df.apply(
            lambda row: f"ID: {row['ID Пятна']} (кандидатов: {row['Кол-во канд.']})", axis=1
        )
        
        selected_option = st.radio(
            "Шаг 1: Выберите инцидент для формирования отчета",
            options=reportable_incidents_df['display_option'].tolist(),
            key="report_selection_radio"
        )

        if selected_option:
            # --- Часть 2: Подготовка данных и генерация отчета ---
            selected_spill_id = selected_option.split(' ')[1]
            spill_row = spills_gdf[spills_gdf['spill_id'] == selected_spill_id].iloc[0]
            candidates_for_spill = candidates_df_for_analytics[candidates_df_for_analytics['spill_id'] == selected_spill_id]
            
            # Генерация графика в памяти
            plot_bytes = create_incident_plot(spill_row, candidates_for_spill)
            
            # Подготовка остальных данных для отчета
            prime_suspect_for_spill = prime_suspects_df[prime_suspects_df['spill_id'] == selected_spill_id] if not prime_suspects_df.empty else pd.DataFrame()
            
            historical_data = {}
            if not prime_suspect_for_spill.empty:
                suspect_mmsi = prime_suspect_for_spill.iloc[0]['mmsi']
                unique_incidents = candidates_df_for_analytics.drop_duplicates(subset=['mmsi', 'spill_id'])
                ship_incident_counts = unique_incidents.groupby('mmsi').size().reset_index(name='incident_count')
                ship_area_sum = unique_incidents.groupby('mmsi')['area_sq_km'].sum().reset_index(name='total_area_sq_km')
                incident_count = ship_incident_counts[ship_incident_counts['mmsi'] == suspect_mmsi]
                area_sum = ship_area_sum[ship_area_sum['mmsi'] == suspect_mmsi]
                historical_data['incident_count'] = incident_count['incident_count'].iloc[0] if not incident_count.empty else 0
                historical_data['total_area_sq_km'] = area_sum['total_area_sq_km'].iloc[0] if not area_sum.empty else 0

            # Генерация самого DOCX файла
            report_bytes = generate_docx_report(spill_row, candidates_for_spill, prime_suspect_for_spill, historical_data, plot_bytes)

            st.markdown("---")
            st.subheader("Шаг 2: Скачайте готовый отчет")
            st.download_button(
                label="✅ Скачать отчет (.docx)",
                data=report_bytes,
                file_name=f"Отчет_{selected_spill_id.replace(':', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    else:
        st.info("Для формирования отчета не найдено инцидентов с судами-кандидатами в заданном временном окне и с учетом фильтров.")